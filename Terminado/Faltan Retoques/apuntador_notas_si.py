from docx import Document
import sys
import time
import subprocess
#                                               Imports
#--------------------------------------------------------------------------------------------------------------
#                                               Variables
doc = Document() #Creamos el Word

nombre = input("¿Cual es tu nombre?: ") #Preguntamos nombre de usuario para ponerlo en el word

ascript = f'''
    set thePath to choose file name with prompt "Seleccione dónde guardar el archivo Word:" default name "notas_de_{nombre}.docx"
    return POSIX path of thePath
''' #Esto es confuso de entender incluso para mi, pero es un script especifico para macOS llamado applescript que hace que salga un menú que te diga donde quieres guardar el archivo word

#----------------------------------------------------------------------------------------------------------------
#                                               Codigo 
print("Bienvenido a el apuntador que apunta los apuntes apuntados por el apuntador.")#Bienvenido a el apuntador que apunta los apuntes apuntador por el apuntador.

time.sleep(2)

print("Escribe las notas que quieras y cuando termines escribe FIN")#Escribe las notas que quieras y cuando termines escribe FIN

time.sleep(1)

print("Recuerda que para cambiar de linea es Enter")#Recuerda que para cambiar de linea es Enter

time.sleep(1)

print("Ya puedes escribir:")#Ya puedes escribir:

lineas = [] #[] Es una lista que almacena las lineas que escribas
while True:
    linea = input() #Aqui es donde tu pones las lineas y lo que quieras escribir en ellas
    if linea == "FIN":#Si pones FIN, acaba la nota y terminas de escribir para llevarlo al proceso de exportación al word
        break
    lineas.append(linea)#Mete en la lista de lineas la linea que acabas de escribir

notas = "\n".join(lineas) #Esto junta TODAS las lineas escritas en un solo bloque
time.sleep(1)

print("Se ha registrado correctamente tus notas")

#---------------------------------------------------------------------------------------------------------------
#                                               Parte Word

destino_archivo = subprocess.check_output(['osascript', '-e', ascript]).decode('utf-8').strip() #Esto es como el script de arriba, es el menú que sale ya llamado para guardar el archivo .word

doc.add_heading(0) #Ponemos los datos del word que se han preguntado anteriormente
if destino_archivo:
    if not destino_archivo.lower().endswith('.docx'):  #Esto es para que tenga la extensión de Word
            destino_archivo += '.docx'
    doc.add_paragraph(notas, style='List Bullet') #Font y añadimos el texto
    doc.save(destino_archivo) #Guardamos
    print(f"Éxito: Archivo guardado en {destino_archivo}") #Confirmación



nombre_archivo = f"Notas de {nombre} " #Guardamos el word como Notas de (y el nombre del usuario)
doc.save(nombre_archivo) #Guardamos el Word

print(f"Se ha exportado correctamente '{nombre_archivo}', Gracias por utilizar mi programa :)") #Fin del programa