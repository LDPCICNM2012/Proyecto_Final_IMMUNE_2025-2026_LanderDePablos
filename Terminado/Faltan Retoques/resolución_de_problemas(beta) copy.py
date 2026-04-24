#AVISO. Para que este codigo funcione, necesitas tener Ollama con el modelo "qwen2.5-coder:3b" instalado. Aparte, necesitarás un mac por los applescript. Lo siento, ya haré un port para windows.
import ollama 
from docx import Document
import subprocess

# --- CONFIGURACIÓN INICIAL ---
conversacion = [] # Aquí guardaremos toda la conversación para que la IA tenga memoria y para exportarla luego a Word.
nombre = input("¿Cual es tu nombre?: ") #Preguntamos el nombre para ponerlo en el Word
doc = Document() #Creamos el word

instrucciones = ( #Ponemos las intrucciones que le vamos a dar a la ia
    "Eres un Solucionador de Problemas de alto rendimiento. "
    "Tu enfoque es: Analizar el problema -> Identificar la causa -> Dar solución técnica/práctica. -> Dar consejos para evitarlo en el futuro. "
    "Responde siempre con estructura de puntos clave y da conversación para acompañar y dar consejos."
)

#Funcion de la IA para responder con memoria y en tiempo real
def proceso_ia_responder(historial_mensajes): #Nombramos a la funcion y le pasamos la memoria de la conversación
    print("\n[IA] Pensando...", end="\r")
    
    # Activamos stream=True para que el texto salga poco a poco
    response = ollama.chat(
        model="llama3.2", #Modelo de ia que vamos a usar
        messages=[
            {'role': 'system', 'content': instrucciones},
            *historial_mensajes # Aquí le pasamos toda la memoria
        ],
        options={"temperature": 0.2, "num_predict": 1000}, #Bajamos la temperatura para que sea más preciso y aumentamos num_predict para que sea más largo, aunque con 1000 ya es suficiente para respuestas detalladas.
        stream=True
    )

    respuesta_completa = ""
    print("[IA]: ", end="") # Etiqueta para la respuesta
    
    for chunk in response: 
        contenido = chunk['message']['content']
        print(contenido, end='', flush=True) # El truco del tiempo real
        respuesta_completa += contenido # Vamos acumulando la respuesta completa para devolverla al final
    
    print("\n") # Espacio al terminar
    return respuesta_completa # Devolvemos la respuesta completa para que se guarde en la memoria y se pueda exportar luego a Word.

#AppleScript para macOS
ascript = f'''
    set thePath to choose file name with prompt "Seleccione dónde guardar el archivo Word:" default name "Conversación de {nombre}.docx"
    return POSIX path of thePath
'''

#BUCLE PRINCIPAL
print(f"Bienvenido {nombre} al Ayudador ayudadoso.") #Presentación
print("Escribe tu pregunta o preocupación.") #Presentación


while True: #Mientras el porgrama esté corriendo:
    pregunta = input(f"{nombre}: ") #Preguntamos el nombre
    
    if pregunta.upper() == "FIN": #Cuando el usuario escriba FIN:
        break #Terminamos el programa
        
    conversacion.append({'role': 'user', 'content': pregunta}) #Se guarda la pregunta del usuario.
    
    res_ia = proceso_ia_responder(conversacion) #Llamamos a la ia para que responda.
    
    conversacion.append({'role': 'assistant', 'content': res_ia}) #Guardamos la respuesta de la ia en la conversación 

#EXPORTACIÓN A WORD
if conversacion: #Si hay conversación para guardar:
    try:
        print("\n--- Guardando conversación ---") #Mensaje de que se esta guardando 
        destino_archivo = subprocess.check_output(['osascript', '-e', ascript]).decode('utf-8').strip() #Ejecutamos el Ascript para que salga el menú de guardado

        if destino_archivo: #Si el usuario elije un destino(que obiamente lo hará):
            if not destino_archivo.lower().endswith('.docx'): #Nos aseguramos de que el archivo tenga la extensión .docx
                destino_archivo += '.docx' #Si no la tiene, se la añadimos automáticamente
            
            doc.add_heading(f'Conversación de {nombre}', 0) #Añadimos un título al Word con el nombre del usuario
            
            for mensaje in conversacion: #Por cada mensaje en la conversación:
                rol = "Tú" if mensaje['role'] == 'user' else "IA" #Ponemos "Tú" para los mensajes del usuario y "IA" para los mensajes de la inteligencia artificial
                p = doc.add_paragraph()#Añadimos un nuevo párrafo para cada mensaje
                p.add_run(f"{rol}: ").bold = True 
                p.add_run(mensaje['content'])#Añadimos el contenido del mensaje al párrafo
            
            doc.save(destino_archivo)#Guardamos el documento en la ruta establecida
            print(f"Éxito: Guardado en {destino_archivo}") #Mensaje de confirmación
    except Exception as e:
        print(f" Error al guardar: {e}") #Si hay error de guardado, que diga el error
else:
    print("No hubo conversación para guardar.") #Si no hay conversación para guardar se lo decimos al usuario.