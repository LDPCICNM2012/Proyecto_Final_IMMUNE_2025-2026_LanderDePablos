#AVISO. Para que este codigo funcione, necesitas tener Ollama con el modelo "qwen2.5-coder:3b" instalado. Aparte, necesitarás un mac por los applescript. Lo siento, ya haré un port para windows.
#Resolución de problemas.py
#Imports
import ollama 
from docx import Document
import subprocess
import time

conversacion = [] #Esto es una lista que guarda la conversación(como la memoria de chatgpt)
run_conversacion = True #Esto es una variable que dice que la conversación esta iniciada.

nombre = input("¿Cual es tu nombre?: ") #Preguntamos nombre de usuario para ponerlo en el word
doc = Document() #Creamos el documento de Word


ascript = f'''
    set thePath to choose file name with prompt "Seleccione dónde guardar el archivo Word:" default name "Apuntes de {nombre}.docx"
    return POSIX path of thePath
''' #Script de Apple necesario para que salga el menú donde puedas elegir donde guardar el Word. Lo de nombre es para que el Word de llame Apuntes de y tu nombre especificado anteriormaente.

print("Bienvenido al Ayudador ayudadoso que ayuda a los necesitados a ayudar sus problemas ayudandose.")#Intro épica
while run_conversacion == True: #Mientras la conversación esta iniciada:_
    pregunta = input("¿Cuál es tu problema? (Escribe 'FIN' para terminar): ")
    if pregunta.upper() == "FIN": #Cuando el usuario escriba "FIN" Termina la conversación
        run_conversacion = False
        break
    print("\n[IA] Generando respuesta...") #Pone en la terminal que la ia esta resumiendo el texto. El \n es para que haga un salto de linea y quede todo ordenadito
    conversacion.append({'role': 'user', 'content': pregunta})
    
    instrucciones = ( #Instrucciones que se le da a la IA para saber como actuar
        "Eres un Solucionador de Problemas de alto rendimiento. "
        "Tu enfoque es: Analizar el problema -> Identificar la causa -> Dar solución técnica/práctica aunque puedes poner comentarios para ser mas resolutivo. "
        "Analiza por qué y da un plan de acción real para la ocasion. "
        "Responde siempre con estructura adecuada al problema diciendo puntos clave."
        "Tambien da un poco de conversación para que el usuario se sienta acompañado, pero sin perder el enfoque de resolver el problema. "
    )
    
    response = ollama.chat(
        model="qwen2.5-coder:1.5b", #El modelo de ia que se va a utilizar
        messages=[
            {'role': 'system', 'content': instrucciones}, #Le damos las instrucciones de antes
            *conversacion
        ],
        options={
            "temperature": 0.3,
            "num_predict": 1000
        }
    )
    
    respuesta_ia = response['message']['content'] #Esta es la respuesta de la ia que guardamos en la variable respuesta_ia
    print(f"[IA] {respuesta_ia}") #Enseñamos la respuesta de
    conversacion.append({'role': 'assistant', 'content': respuesta_ia}) #Guardamos la respuesta de la ia en la lista de conversacion.

#Parte de Word.
try: #El try es para que haga lo que hay abajo y si no va que no pete como un tonto si no que te diga pq. 
    print("Abriendo menú para que guardes tu archivo...")
    destino_archivo = subprocess.check_output(['osascript', '-e', ascript]).decode('utf-8').strip() #Aqui llamamos al applescript para que salga el menú.

    if destino_archivo:
        if not destino_archivo.lower().endswith('.docx'):#Basicamente (que esto no lo hago ni yo) es para que el usuario si no pone la terminacion de word .docx la pone automaticamente. Si no sería horrible.
            destino_archivo += '.docx' #La extensión.
        
        doc.add_heading(f'Conversación de {nombre}', 0) #Pone titulo al Word con el nombre del usuario.
        for mensaje in conversacion:
                rol = "Tú" if mensaje['role'] == 'user' else "IA"
                # Añadimos un párrafo por cada intervención
                p = doc.add_paragraph()
                p.add_run(f"{rol}: ").bold = True # Nombre en negrita
                p.add_run(mensaje['content']) # Contenido en texto normal
        
        doc.save(destino_archivo) #Destino del archivo que se guardará donde haya escogido el usuario.
        print(f"\n Éxito: Archivo guardado en {destino_archivo}")#Confirmación de que se ha guardado el archivo.")
except Exception as e:
    print(f"Error al guardar: {e}") #Si hay un error, que te diga pq y donde.