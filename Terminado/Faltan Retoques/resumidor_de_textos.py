#AVISO. Para que este codigo funcione, necesitas tener Ollama con el modelo "qwen2.5-coder:3b" instalado. Aparte, necesitarás un mac por los applescript. Lo siento, ya haré un port para windows.
#IMPORTS
import ollama
from docx import Document
import subprocess
import time

#CONFIGURACIÓN DE LA IA
def proceso_ia(texto): #Funcion que hace que la ia procese los datos que se han puesto por el usuario con x instrucciones. Basicamente lo que engloba todo lo que se encarga la ia
    print("\n[IA] Resumiendo texto...") #Pone en la terminal que la ia esta resumiendo el texto. El \n es para que haga un salto de linea y quede todo ordenadito
    
    #Intrucciones predeterminadas a la ia
    instrucciones = (
        "Eres un experto en el tema proporcionado. Tu conocimiento se basa estrictamente en hechos reales. "
        "REGLA DE SEGURIDAD ABSOLUTA: Solo puedes responder a temas que pertenezcan al ámbito educativo, "
        "académico, histórico o laboral. Si el usuario te pide algo fuera de estos ámbitos, DEBES responder "
        "ÚNICAMENTE con la frase: 'ERROR: La petición no pertenece al ámbito educativo o laboral.' "
        "Si la petición es válida, redacta un texto muy extenso, preciso y con párrafos bien estructurados "
        "explicando el contexto, las causas y las consecuencias. No inventes datos bajo ninguna circunstancia." \
        "Si el usuario te pasa algun archivo para que resumas, analiza que es para fines legales y resumelo, si no di que es error (Si es legal pero no es educativo resumelo)"
    )

    response = ollama.chat(
        model="llama3.2",
        messages=[
            {'role': 'system', 'content': instrucciones},
            {'role': 'user', 'content': f"Desarrolla o resume de manera extensa y rigurosa el siguiente tema/apuntes: {texto}"}
        ],
        options={
            "temperature": 0.2,  # Bajamos la temperatura a 0.2 para que sea preciso, riguroso y no invente datos.
            "num_predict": 2000  # Aumentado a 2000 para que el texto sea aún más largo y detallado.
        }
    )
    return response['message']['content'] #"Return" significa devuelve, a si que esto lo que hace es devolver lo que da la ia

# Variables y AppleScript
doc = Document() #Creamos el documento de Word
nombre = input("¿Cual es tu nombre?: ") #Te pregunta el nombre para ponerlo en el Word y para ser educados :)

ascript = f'''
    set thePath to choose file name with prompt "Seleccione dónde guardar el archivo Word:" default name "Apuntes de {nombre}.docx"
    return POSIX path of thePath
''' #Script de Apple necesario para que salga el menú donde puedas elegir donde guardar el Word. Lo de nombre es para que el Word de llame Apuntes de y tu nombre especificado anteriormaente.

#Presentacion y recogida de datos

print("Bienvenido al Lander_Resumidor 2.0 Pro Plus Ultra Deluxe Max Edition (Only for Mac. Only for personal use. Copyright 2026 Lander S.L and Immune Tecnology Institute. All rights reserved.)")#Presentacion muy presentada para que quede formal. Todo lo escritø es necesario.
time.sleep(1)
print("Escribe tus Apuntes o tema que quieres resumir. Escribe 'FIN' en una línea nueva para terminar:") #Aqui esta lo de mi codigo para que el usuario escriba sus nota/instrucciones y se lo mande a la ia sin limite de lineas.
lineas = [] #Lista donde se guardan las lineas que el usuario va escribiendo
while True:
    linea = input("> ") #Esto es para que salga ete logo > y sepas que tienes que escribir ahí, detallitos importantes. 
    if linea.upper() == "FIN": #Si pones fin en mayusculas sales de escribir a la ia y se manda
        break
    lineas.append(linea)

Apuntes_puras = "\n".join(lineas) #Copila todas las lineas y lo hace en un solo bloque para mandarselo a la ia. 

Apuntes_finales = proceso_ia(Apuntes_puras) #Esto es lo bueno, integración de ia. Esto es basicamente el "promp" que se manda a la ia. Es como si fueras a chatgpt y le mandaras lo que has escrito + todas las instrucciones de arriba pero mas simplificado y user-friendly.

#Parte de Word.
try: #El try es para que haga lo que hay abajo y si no va que no pete como un tonto si no que te diga pq. 
    print("Abriendo menú para que guardes tu archivo...")
    destino_archivo = subprocess.check_output(['osascript', '-e', ascript]).decode('utf-8').strip() #Aqui llamamos al applescript para que salga el menú.

    if destino_archivo:
        if not destino_archivo.lower().endswith('.docx'):#Basicamente (que esto no lo hago ni yo) es para que el usuario si no pone la terminacion de word .docx la pone automaticamente. Si no sería horrible.
            destino_archivo += '.docx' #La extensión.
        
        doc.add_heading(f'Apuntes de {nombre}', 0) #Pone titulo al Word con el nombre del usuario.
        doc.add_paragraph(Apuntes_finales)#Pone los Apuntes ya procesados por la ia en el Word.
        
        doc.save(destino_archivo) #Destino del archivo que se guardará donde haya escogido el usuario.
        print(f"\n Éxito: Archivo guardado en {destino_archivo}")#Confirmación de que se ha guardado el archivo.")
except Exception as e:
    print(f"Error al guardar: {e}") #Si hay un error, que te diga pq y donde.