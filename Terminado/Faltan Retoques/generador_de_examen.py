#AVISO. Para que este codigo funcione, necesitas tener Ollama con el modelo "qwen2.5-coder:3b" instalado. Aparte, necesitarás un mac por los applescript. Lo siento, ya haré un port para windows.
#IMPORTS
from httpcore import stream
import ollama
from docx import Document
import subprocess
import time

#-----------------------------------------------FUNCIONES-------------------------------------------------
#INSTRUCCIONES PARA LA IA SOBRE CÓMO DEBE SER EL EXAMEN Y LAS EXPORTE A WORD
instrucciones1 = """El examen debe tener:
        1. Un título que tenga que ver con el tema.
        2. El numero de preguntas que elija el usuario de opción múltiple (A, B, C, D, E) o rellena la respuesta dejando un hueco en blanco.
        3. La mitad de las preguntas que haya dicho el usuario que sean de desarrollo (explicar conceptos)(Si el numero no es par, se redondea hacia arriba).
        4. Cuando el usuario responda a las preguntas, debes evaluar sus respuestar sobre 10 dandole la calificacion y explicando porque el fallo o el acierto.
        5. Si el usuario dice que otro examen, generas otro examen diferente al anterior pero con las mismas características. Si el usuario dice que no, le dices que hasta luego y se acaba el programa."""

#INSTRUCCIONES PARA LA IA SOBRE CÓMO DEBE SER EL EXAMEN Y LAS RESPUESTA EN EL PROGRAMA

instrucciones2 = """"Eres un evaluador académico profesional. Tu tarea es crear un examen sobre el tema proporcionado."
        "El examen debe tener:"
        "1. Un título que tenga que ver con el tema."
        "2. El numero de preguntas que elija el usuario de opción múltiple (A, B, C, D, E) o rellena la respuesta dejando un hueco en blanco."
        "3. La mitad de las preguntas que haya dicho el usuario que sean de desarrollo (explicar conceptos)(Si el numero no es par, se redondea hacia arriba)."
        "4. Cuando el usuario responda a las preguntas, debes evaluar sus respuestar sobre 10 dandole la calificacion y explicando porque el fallo o el acierto."
        "5. Si el usuario dice que otro examen, generas otro examen diferente al anterior pero con las mismas características. Si el usuario dice que no, le dices que hasta luego y se acaba el programa."
        "Solo pon la explicación y la respuesta cuando el usuario haya respondido a todas las preguntas, no antes."
        "NO PONGAS LAS RESPUESTAS HASTA QUE EL USUARIO TE DIGA LAS RESPUESTAR (EJ. 1.A 2.C 3.B etc...)"""
#CONFIGURACIÓN DE LA IA PARA QUE EL USUARIO ELIJA EXPORTAR A WORD

def proceso_ia_exportar(texto): #Funcion que hace que la ia procese los datos que se han puesto por el usuario con x instrucciones. Basicamente lo que engloba todo lo que se encarga la ia
    print("\n[IA] Generando examen...") #Pone en la terminal que la ia esta resumiendo el texto. El \n es para que haga un salto de linea y quede todo ordenadito
    
    #Intrucciones predeterminadas a la ia
    instrucciones = (instrucciones1)

    response = ollama.chat(
        model="llama3.2",
        messages=[
            {'role': 'system', 'content': instrucciones},
            {'role': 'user', 'content': f"Hazme un examen sobre: {texto} con las características mencionadas anteriormente."}
        ],
        stream=True,
        options={
            "temperature": 0.2,  # Bajamos la temperatura a 0.2 para que sea preciso, riguroso y no invente datos.
            "num_predict": 2000  # Aumentado a 2000 para que el texto sea aún más largo y detallado.
        }
    )

    respuesta_completa = ""
    for chunk in response: 
        contenido = chunk['message']['content']
        print(contenido, end='', flush=True) # El truco del tiempo real
        respuesta_completa += contenido # Vamos acumulando la respuesta completa para devolverla al final
    print("\n") # Espacio al terminar
    return respuesta_completa # Devolvemos la respuesta completa para que se guarde en la memoria y se pueda exportar luego a Word.


#CONFIGURACIÓN DE LA IA PARA QUE EL USUARIO ELIJA RESPONDER EN EL POROGRAMA

def proceso_ia_responder(texto): #Funcion que hace que la ia procese los datos que se han puesto por el usuario con x instrucciones. Basicamente lo que engloba todo lo que se encarga la ia
    print("\n[IA] Generando examen...") #Pone en la terminal que la ia esta resumiendo el texto. El \n es para que haga un salto de linea y quede todo ordenadito
    
    #Intrucciones predeterminadas a la ia
    instrucciones = (instrucciones2)


    response = ollama.chat(
        model="llama3.2",
        messages=[
            {'role': 'system', 'content': instrucciones},
            {'role': 'user', 'content': f"Hazme un examen sobre: {texto} con las características mencionadas anteriormente."}
        ],
        stream=True,
        options={
            "temperature": 0.2,  # Bajamos la temperatura a 0.2 para que sea preciso, riguroso y no invente datos.
            "num_predict": 2000  # Aumentado a 2000 para que el texto sea aún más largo y detallado.
        }
    )
    respuesta_completa = ""
    for chunk in response: 
        contenido = chunk['message']['content']
        print(contenido, end='', flush=True) # El truco del tiempo real
        respuesta_completa += contenido # Vamos acumulando la respuesta completa para devolverla al final
    print("\n") # Espacio al terminar
    return respuesta_completa # Devolvemos la respuesta completa para que se guarde en la memoria y se pueda exportar luego a Word.

#-----------------------------------------------PROGRAMA PRINCIPAL-------------------------------------------------

# Variables y AppleScript
while True:
    doc = Document() #Creamos el documento de Word
    nombre = input("¿Cual es tu nombre?: ") #Te pregunta el nombre para ponerlo en el Word y para ser educados :)

    ascript = f'''
        set thePath to choose file name with prompt "Seleccione dónde guardar el archivo Word:" default name "Apuntes de {nombre}.docx"
        return POSIX path of thePath
    ''' #Script de Apple necesario para que salga el menú donde puedas elegir donde guardar el Word. Lo de nombre es para que el Word de llame Apuntes de y tu nombre especificado anteriormaente.

    #Presentacion y recogida de datos

    print("Bienvenido al Lander_Creador_De_Examenes 2.0 Pro Plus Ultra Deluxe Max Edition (Only for Mac. Only for personal use. Copyright 2026 Lander S.L and Immune Tecnology Institute. All rights reserved.)")#Presentacion muy presentada para que quede formal. Todo lo escritø es necesario.
    time.sleep(1)
    print("Escribe el tema sobre el que quieres hacer el examen. Escribe 'FIN' en una línea nueva para terminar:") #Aqui esta lo de mi codigo para que el usuario escriba sus nota/instrucciones y se lo mande a la ia sin limite de lineas.
    lineas = [] #Lista donde se guardan las lineas que el usuario va escribiendo
    while True:
        linea = input("> ") #Esto es para que salga ete logo > y sepas que tienes que escribir ahí, detallitos importantes. 
        if linea.upper() == "FIN": #Si pones fin en mayusculas sales de escribir a la ia y se manda
            break
        lineas.append(linea)

    Apuntes_puras = "\n".join(lineas) #Copila todas las lineas y lo hace en un solo bloque para mandarselo a la ia. 

    #----------------------------------------------Condicionales-----------------------------------------------------------
    destino_examen = input("¿Que prefieres?: 1. Exportar el examen a Word y responderlo ahí(luego tendrás que copiar la ruta del archivo y pasarla al otro proyecto) o 2. Responder aquí mismo en el programa? (Escribe 1 o 2): ") #Aqui le das al usuario la opción de exportar a Word o responder en el programa.

    if destino_examen == "1":
        Apuntes_finales = proceso_ia_exportar(Apuntes_puras) #Esto es lo bueno, integración de ia. Esto es basicamente el "promp" que se manda a la ia. Es como si fueras a chatgpt y le mandaras lo que has escrito + todas las instrucciones de arriba pero mas simplificado y user-friendly.
        try: #El try es para que haga lo que hay abajo y si no va que no pete como un tonto si no que te diga pq. 
            print("Abriendo menú para que guardes tu archivo...")
            destino_archivo = subprocess.check_output(['osascript', '-e', ascript]).decode('utf-8').strip() #Aqui llamamos al applescript para que salga el menú.

            if destino_archivo:
                if not destino_archivo.lower().endswith('.docx'):#Basicamente (que esto no lo hago ni yo) es para que el usuario si no pone la terminacion de word .docx la pone automaticamente. Si no sería horrible.
                    destino_archivo += '.docx' #La extensión.
                
                doc.add_heading(f'Examen de {nombre}', 0) #Pone titulo al Word con el nombre del usuario.
                doc.add_paragraph(Apuntes_finales)#Pone los Apuntes ya procesados por la ia en el Word.
                
                doc.save(destino_archivo) #Destino del archivo que se guardará donde haya escogido el usuario.
                print(f"\n Éxito: Archivo guardado en {destino_archivo}")#Confirmación de que se ha guardado el archivo.")
        except Exception as e:
            print(f"Error al guardar: {e}") #Si hay un error, que te diga pq y donde.

    elif destino_examen == "2":
        Apuntes_finales = proceso_ia_responder(Apuntes_puras) 
        print("\n[IA] Aquí está tu examen. Recuerda, el objetivo es aprobar, no suspender:\n")
        
        historial = [
            {"role": "system", "content": instrucciones2},
            {"role": "assistant", "content": Apuntes_finales}
        ]

        while True:
            respuesta_usuario = input(f"\n{nombre} (Escribe tu respuesta o 'SALIR' para ir al menú): ")
            
            if respuesta_usuario.upper() == "SALIR":
                break
            
            historial.append({"role": "user", "content": respuesta_usuario})
            print("\n[IA] Corrigiendo...\n")
            
            # Llamada con stream para ver la corrección en tiempo real
            stream_correcion = ollama.chat(model="llama3.2", messages=historial, stream=True)
            
            comentario_ia = ""
            for chunk in stream_correcion:
                contenido = chunk['message']['content']
                print(contenido, end='', flush=True)
                comentario_ia += contenido
            
            print("\n") # Salto de línea al terminar la corrección
            
            # Guardamos el comentario en el historial para que la IA recuerde lo que dijo
            historial.append({"role": "assistant", "content": comentario_ia})


#COSAS QUE FALTAN:
#-QUE CORRIGA BIEN EL CBRN
#-PROBAR EL EXPORTE A WORD 